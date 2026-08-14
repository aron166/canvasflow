"""
Source ingestion — turns a URL or uploaded file into plain text.

Every resolver returns a `ResolvedSource`: the extracted text plus enough metadata for
the UI to show what was ingested and for retrieval to cite it later (timestamps for
transcripts, page numbers for PDFs).

Heavy dependencies are imported lazily inside their resolver, so a missing
`faster-whisper` only breaks media transcription — the rest of the app still runs.
"""

from __future__ import annotations

import asyncio
import os
import re
import subprocess
from dataclasses import dataclass, field
from enum import Enum
from typing import Any
from urllib.parse import parse_qs, urlparse

import httpx

MAX_DOWNLOAD_BYTES = 100 * 1024 * 1024  # 100 MB — a generous ceiling for a local tool
FETCH_TIMEOUT_SECONDS = 60


class SourceKind(str, Enum):
    YOUTUBE = "youtube"
    WEBPAGE = "webpage"
    DOCUMENT = "document"  # PDF / DOCX / TXT / MD
    MEDIA = "media"  # local audio or video, transcribed
    TEXT = "text"  # pasted directly


class IngestError(RuntimeError):
    """A failure worth showing verbatim on the node card — so it must say what to do."""


@dataclass
class Segment:
    """A citable span of the source. `label` is what a citation renders as."""

    text: str
    label: str  # e.g. "12:04" or "p. 7"
    start: float | None = None  # seconds, transcripts only


@dataclass
class ResolvedSource:
    kind: SourceKind
    title: str
    text: str
    origin: str  # URL or filename
    segments: list[Segment] = field(default_factory=list)
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text)


# --------------------------------------------------------------------------------------
# URL classification
# --------------------------------------------------------------------------------------

_YOUTUBE_HOSTS = {
    "youtube.com",
    "www.youtube.com",
    "m.youtube.com",
    "music.youtube.com",
    "youtu.be",
    "www.youtu.be",
}


def youtube_video_id(url: str) -> str | None:
    """Extracts the 11-character video ID from any common YouTube URL shape."""
    try:
        parsed = urlparse(url if "//" in url else f"https://{url}")
    except ValueError:
        return None

    host = (parsed.hostname or "").lower()
    if host not in _YOUTUBE_HOSTS:
        return None

    # youtu.be/<id>
    if host.endswith("youtu.be"):
        candidate = parsed.path.lstrip("/").split("/")[0]
        return candidate if _valid_video_id(candidate) else None

    # /watch?v=<id>
    if parsed.path == "/watch":
        candidate = parse_qs(parsed.query).get("v", [""])[0]
        return candidate if _valid_video_id(candidate) else None

    # /shorts/<id>, /embed/<id>, /live/<id>
    match = re.match(r"^/(?:shorts|embed|live|v)/([^/?#]+)", parsed.path)
    if match and _valid_video_id(match.group(1)):
        return match.group(1)

    return None


def _valid_video_id(candidate: str) -> bool:
    return bool(re.fullmatch(r"[A-Za-z0-9_-]{11}", candidate))


def classify_url(url: str) -> SourceKind:
    if youtube_video_id(url):
        return SourceKind.YOUTUBE
    path = urlparse(url).path.lower()
    if path.endswith((".pdf", ".docx", ".txt", ".md")):
        return SourceKind.DOCUMENT
    if path.endswith((".mp3", ".mp4", ".wav", ".m4a", ".mov", ".webm", ".mkv", ".ogg")):
        return SourceKind.MEDIA
    return SourceKind.WEBPAGE


def timestamp_label(seconds: float) -> str:
    total = int(seconds)
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours}:{minutes:02d}:{secs:02d}" if hours else f"{minutes}:{secs:02d}"


# --------------------------------------------------------------------------------------
# YouTube
# --------------------------------------------------------------------------------------


async def resolve_youtube(url: str) -> ResolvedSource:
    """Pulls the published transcript. No download, no API key, no quota.

    Falls back to transcribing the audio only if the video has no transcript at all —
    which is why this resolver can be slow on the rare uncaptioned video.
    """
    video_id = youtube_video_id(url)
    if not video_id:
        raise IngestError(f"'{url}' is not a recognisable YouTube URL.")

    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        from youtube_transcript_api._errors import (
            NoTranscriptFound,
            TranscriptsDisabled,
            VideoUnavailable,
        )
    except ImportError as exc:
        raise IngestError(
            "youtube-transcript-api is not installed. Run: pip install -r requirements.txt"
        ) from exc

    def _fetch() -> list[dict[str, Any]]:
        api = YouTubeTranscriptApi()
        # Prefer a real transcript; fall back to auto-generated captions.
        fetched = api.fetch(video_id, languages=["en", "en-US", "en-GB"])
        return [
            {"text": snippet.text, "start": snippet.start, "duration": snippet.duration}
            for snippet in fetched.snippets
        ]

    try:
        raw = await asyncio.to_thread(_fetch)
    except VideoUnavailable as exc:
        raise IngestError("That video is unavailable (private, deleted, or region-locked).") from exc
    except (TranscriptsDisabled, NoTranscriptFound):
        # No published captions. Fall back to downloading the audio and transcribing it
        # locally — far slower, so it is never the first choice.
        return await _transcribe_youtube_audio(video_id)
    except Exception as exc:  # noqa: BLE001 - the library raises a wide range
        raise IngestError(f"Could not fetch the transcript: {type(exc).__name__}: {exc}") from exc

    if not raw:
        raise IngestError("The transcript came back empty.")

    segments = [
        Segment(
            text=entry["text"].strip(),
            label=timestamp_label(entry["start"]),
            start=float(entry["start"]),
        )
        for entry in raw
        if entry.get("text", "").strip()
    ]

    # Timestamps are inlined so the model can cite a moment, not just a claim.
    text = "\n".join(f"[{segment.label}] {segment.text}" for segment in segments)
    duration = segments[-1].start if segments and segments[-1].start else 0.0

    return ResolvedSource(
        kind=SourceKind.YOUTUBE,
        title=await _youtube_title(video_id) or f"YouTube {video_id}",
        text=text,
        origin=f"https://www.youtube.com/watch?v={video_id}",
        segments=segments,
        meta={"video_id": video_id, "duration_label": timestamp_label(duration)},
    )


async def _transcribe_youtube_audio(video_id: str) -> ResolvedSource:
    """Last resort for uncaptioned videos: download audio with yt-dlp, run Whisper.

    Costs minutes of CPU where the transcript API costs milliseconds, so this only runs
    when YouTube publishes no captions at all.
    """
    import tempfile

    try:
        import yt_dlp  # noqa: F401
    except ImportError as exc:
        raise IngestError(
            "This video has no published transcript. Transcribing it locally needs "
            "yt-dlp: pip install yt-dlp faster-whisper"
        ) from exc

    if not _has_ffmpeg():
        raise IngestError(
            "This video has no published transcript, and local transcription needs "
            "ffmpeg. Install ffmpeg and retry."
        )

    with tempfile.TemporaryDirectory() as workdir:
        target = os.path.join(workdir, "audio.%(ext)s")

        def _download() -> str:
            import yt_dlp

            options = {
                "format": "bestaudio/best",
                "outtmpl": target,
                "quiet": True,
                "no_warnings": True,
                "noprogress": True,
                # Whisper wants mono 16k; ffmpeg does the conversion on extraction.
                "postprocessors": [
                    {
                        "key": "FFmpegExtractAudio",
                        "preferredcodec": "mp3",
                        "preferredquality": "128",
                    }
                ],
            }
            with yt_dlp.YoutubeDL(options) as downloader:
                downloader.download([f"https://www.youtube.com/watch?v={video_id}"])

            for name in os.listdir(workdir):
                if name.startswith("audio."):
                    return os.path.join(workdir, name)
            raise IngestError("yt-dlp finished but produced no audio file.")

        try:
            audio_path = await asyncio.to_thread(_download)
        except IngestError:
            raise
        except Exception as exc:  # noqa: BLE001 - yt-dlp raises many shapes
            raise IngestError(f"Could not download the audio: {type(exc).__name__}: {exc}") from exc

        resolved = await resolve_media(audio_path, f"YouTube {video_id}")

    resolved.kind = SourceKind.YOUTUBE
    resolved.title = await _youtube_title(video_id) or f"YouTube {video_id}"
    resolved.origin = f"https://www.youtube.com/watch?v={video_id}"
    resolved.meta["video_id"] = video_id
    resolved.meta["transcript_source"] = "whisper"  # vs YouTube's own captions
    return resolved


async def _youtube_title(video_id: str) -> str | None:
    """Best-effort title via oEmbed — a missing title must never fail the ingest."""
    try:
        async with httpx.AsyncClient(timeout=15) as client:
            response = await client.get(
                "https://www.youtube.com/oembed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"},
            )
            if response.status_code == 200:
                return str(response.json().get("title") or "").strip() or None
    except Exception:  # noqa: BLE001
        pass
    return None


# --------------------------------------------------------------------------------------
# Web pages
# --------------------------------------------------------------------------------------


async def resolve_webpage(url: str) -> ResolvedSource:
    """Fetches a page and strips nav/ads/boilerplate down to the article body."""
    try:
        import trafilatura
    except ImportError as exc:
        raise IngestError(
            "trafilatura is not installed. Run: pip install -r requirements.txt"
        ) from exc

    try:
        async with httpx.AsyncClient(
            timeout=FETCH_TIMEOUT_SECONDS,
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; CanvasFlow/0.1)"},
        ) as client:
            response = await client.get(url)
    except httpx.HTTPError as exc:
        raise IngestError(f"Could not fetch {url}: {type(exc).__name__}") from exc

    if response.status_code >= 400:
        raise IngestError(f"{url} returned HTTP {response.status_code}.")

    html = response.text
    extracted = await asyncio.to_thread(
        trafilatura.extract,
        html,
        include_comments=False,
        include_tables=True,
        favor_precision=True,
    )

    if not extracted or len(extracted.strip()) < 50:
        raise IngestError(
            "Almost no readable text on that page — it may be JavaScript-rendered or "
            "behind a paywall."
        )

    metadata = await asyncio.to_thread(trafilatura.extract_metadata, html)
    title = (getattr(metadata, "title", None) or urlparse(url).netloc or url).strip()

    return ResolvedSource(
        kind=SourceKind.WEBPAGE,
        title=title,
        text=extracted.strip(),
        origin=url,
        meta={"site": urlparse(url).netloc},
    )


# --------------------------------------------------------------------------------------
# Documents (PDF / DOCX / plain text)
# --------------------------------------------------------------------------------------


async def resolve_document(path: str, display_name: str | None = None) -> ResolvedSource:
    name = display_name or os.path.basename(path)
    suffix = os.path.splitext(path)[1].lower()

    if suffix == ".pdf":
        pages = await asyncio.to_thread(_extract_pdf, path)
    elif suffix == ".docx":
        pages = await asyncio.to_thread(_extract_docx, path)
    elif suffix in {".txt", ".md", ".markdown", ".csv"}:
        pages = await asyncio.to_thread(_extract_plaintext, path)
    else:
        raise IngestError(
            f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, .txt, .md, .csv"
        )

    segments = [
        Segment(text=body, label=label) for label, body in pages if body.strip()
    ]
    if not segments:
        raise IngestError(
            f"No extractable text in {name}. If it's a scanned PDF, it needs OCR first."
        )

    text = "\n\n".join(f"[{segment.label}]\n{segment.text}" for segment in segments)
    return ResolvedSource(
        kind=SourceKind.DOCUMENT,
        title=name,
        text=text,
        origin=name,
        segments=segments,
        meta={"pages": len(segments), "format": suffix.lstrip(".")},
    )


def _extract_pdf(path: str) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise IngestError("pypdf is not installed. Run: pip install -r requirements.txt") from exc

    try:
        reader = PdfReader(path)
    except Exception as exc:  # noqa: BLE001 - pypdf raises many shapes on bad files
        raise IngestError(f"Could not open the PDF: {exc}") from exc

    if reader.is_encrypted:
        # An empty user password is common and decrypts silently; a real one can't be guessed.
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise IngestError("That PDF is password-protected.") from exc

    return [
        (f"p. {index}", (page.extract_text() or "").strip())
        for index, page in enumerate(reader.pages, start=1)
    ]


def _extract_docx(path: str) -> list[tuple[str, str]]:
    try:
        import docx
    except ImportError as exc:
        raise IngestError(
            "python-docx is not installed. Run: pip install -r requirements.txt"
        ) from exc

    document = docx.Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = [
        " | ".join(cell.text.strip() for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    body = "\n".join(paragraphs + tables)
    return [("document", body)]


def _extract_plaintext(path: str) -> list[tuple[str, str]]:
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        return [("document", handle.read())]


# --------------------------------------------------------------------------------------
# Local audio / video
# --------------------------------------------------------------------------------------


def _has_ffmpeg() -> bool:
    from shutil import which

    return which("ffmpeg") is not None


async def resolve_media(
    path: str, display_name: str | None = None, model_size: str = "base"
) -> ResolvedSource:
    """Transcribes local audio/video with faster-whisper. Fully offline."""
    name = display_name or os.path.basename(path)

    if not _has_ffmpeg():
        raise IngestError(
            "ffmpeg is required to decode media. Install it (`sudo dnf install ffmpeg` "
            "or `brew install ffmpeg`) and retry."
        )

    try:
        from faster_whisper import WhisperModel
    except ImportError as exc:
        raise IngestError(
            "faster-whisper is not installed. Run: pip install faster-whisper\n"
            "(It is optional — the other source types work without it.)"
        ) from exc

    def _transcribe() -> list[Segment]:
        # int8 on CPU keeps a 12-core / 7 GB box comfortable; float16 needs a GPU.
        model = WhisperModel(model_size, device="cpu", compute_type="int8")
        segments, _info = model.transcribe(path, vad_filter=True)
        return [
            Segment(
                text=segment.text.strip(),
                label=timestamp_label(segment.start),
                start=float(segment.start),
            )
            for segment in segments
            if segment.text.strip()
        ]

    try:
        segments = await asyncio.to_thread(_transcribe)
    except IngestError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise IngestError(f"Transcription failed: {type(exc).__name__}: {exc}") from exc

    if not segments:
        raise IngestError(f"No speech detected in {name}.")

    text = "\n".join(f"[{segment.label}] {segment.text}" for segment in segments)
    duration = segments[-1].start or 0.0

    return ResolvedSource(
        kind=SourceKind.MEDIA,
        title=name,
        text=text,
        origin=name,
        segments=segments,
        meta={"duration_label": timestamp_label(duration), "whisper_model": model_size},
    )


def probe_media_duration(path: str) -> float | None:
    """Duration via ffprobe, for the UI. None when it can't be determined."""
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "error", "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1", path,
            ],
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        return float(result.stdout.strip()) if result.returncode == 0 else None
    except (OSError, ValueError, subprocess.TimeoutExpired):
        return None


# --------------------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------------------


async def resolve_url(url: str) -> ResolvedSource:
    """Routes a URL to the right resolver."""
    url = url.strip()
    if not url:
        raise IngestError("No URL provided.")
    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"

    kind = classify_url(url)
    if kind is SourceKind.YOUTUBE:
        return await resolve_youtube(url)
    if kind is SourceKind.DOCUMENT:
        return await _resolve_remote_binary(url, SourceKind.DOCUMENT)
    if kind is SourceKind.MEDIA:
        return await _resolve_remote_binary(url, SourceKind.MEDIA)
    return await resolve_webpage(url)


async def _resolve_remote_binary(url: str, kind: SourceKind) -> ResolvedSource:
    """Downloads a linked file to a temp path, then hands it to the local resolver."""
    import tempfile

    suffix = os.path.splitext(urlparse(url).path)[1] or ".bin"
    filename = os.path.basename(urlparse(url).path) or f"download{suffix}"

    try:
        async with httpx.AsyncClient(
            timeout=FETCH_TIMEOUT_SECONDS, follow_redirects=True
        ) as client:
            async with client.stream("GET", url) as response:
                if response.status_code >= 400:
                    raise IngestError(f"{url} returned HTTP {response.status_code}.")
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
                    written = 0
                    async for chunk in response.aiter_bytes():
                        written += len(chunk)
                        if written > MAX_DOWNLOAD_BYTES:
                            raise IngestError(
                                f"{filename} exceeds the {MAX_DOWNLOAD_BYTES // 1024 // 1024} MB "
                                "download limit."
                            )
                        handle.write(chunk)
                    temp_path = handle.name
    except httpx.HTTPError as exc:
        raise IngestError(f"Could not download {url}: {type(exc).__name__}") from exc

    try:
        resolver = resolve_document if kind is SourceKind.DOCUMENT else resolve_media
        resolved = await resolver(temp_path, filename)
        resolved.origin = url
        return resolved
    finally:
        try:
            os.unlink(temp_path)
        except OSError:
            pass


async def resolve_file(path: str, display_name: str | None = None) -> ResolvedSource:
    """Routes an uploaded file by extension."""
    suffix = os.path.splitext(path)[1].lower()
    media_suffixes = {".mp3", ".mp4", ".wav", ".m4a", ".mov", ".webm", ".mkv", ".ogg", ".flac"}
    if suffix in media_suffixes:
        return await resolve_media(path, display_name)
    return await resolve_document(path, display_name)


# --------------------------------------------------------------------------------------
# Self-check:  python ingest.py
# --------------------------------------------------------------------------------------


def _demo() -> None:
    # Every YouTube URL shape must yield the same ID.
    for url in (
        "https://www.youtube.com/watch?v=dQw4w9WgXcQ",
        "https://youtu.be/dQw4w9WgXcQ",
        "https://www.youtube.com/shorts/dQw4w9WgXcQ",
        "https://www.youtube.com/embed/dQw4w9WgXcQ?start=30",
        "m.youtube.com/watch?v=dQw4w9WgXcQ&list=PLxyz",
    ):
        assert youtube_video_id(url) == "dQw4w9WgXcQ", url

    # Non-YouTube and malformed IDs must not be misread as videos.
    assert youtube_video_id("https://example.com/watch?v=dQw4w9WgXcQ") is None
    assert youtube_video_id("https://www.youtube.com/watch?v=tooshort") is None
    assert youtube_video_id("https://vimeo.com/12345") is None

    assert classify_url("https://youtu.be/dQw4w9WgXcQ") is SourceKind.YOUTUBE
    assert classify_url("https://example.com/report.pdf") is SourceKind.DOCUMENT
    assert classify_url("https://example.com/talk.mp4") is SourceKind.MEDIA
    assert classify_url("https://example.com/blog/post") is SourceKind.WEBPAGE

    assert timestamp_label(0) == "0:00"
    assert timestamp_label(75) == "1:15"
    assert timestamp_label(3675) == "1:01:15"

    print("ingest self-check passed")


if __name__ == "__main__":
    _demo()
